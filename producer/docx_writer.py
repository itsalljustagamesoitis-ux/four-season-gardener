"""
Write article review files as .docx.
Uses proper Word heading styles so it looks clean in Word and Pages.
Reads edited .docx back to markdown for the publish step.
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError("Run: pip3 install python-docx")


# Brand colours
GREEN = RGBColor(0x2D, 0x50, 0x16)
GOLD = RGBColor(0xC1, 0x9A, 0x4B)
GREY = RGBColor(0x88, 0x88, 0x88)


def _set_para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def build_docx(article: dict, body: str, title: str, description: str) -> Document:
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Meta block ───────────────────────────────────────────────────────────
    def meta_label(label, value, colour=None):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = GREY
        run_val = p.add_run(value)
        run_val.font.size = Pt(10)
        if colour:
            run_val.font.color.rgb = colour
        _set_para_spacing(p, after=2)
        return p

    meta_label("TITLE", title, GREEN)
    meta_label("DESC", description)

    products_str = ", ".join(article.get("products", [])) or "none"
    hub_str = f"{article.get('hub_label','')} ({article.get('hub_url','')})"
    meta_label("Products", products_str)
    meta_label("Hub", hub_str)
    meta_label(
        "Info",
        f"ID {article['id']} | {article['type']} | KD {article.get('kd',0)} | Vol {article.get('volume',0)}"
    )

    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 60)
    run.font.color.rgb = GREY
    run.font.size = Pt(9)

    # ── Body ─────────────────────────────────────────────────────────────────
    # Parse markdown line by line into Word elements
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # H2
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.style.font.color.rgb = GREEN
            _set_para_spacing(h, before=14, after=4)

        # H3
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            _set_para_spacing(h, before=8, after=3)

        # Bullet list item
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:].strip())
            _set_para_spacing(p, after=3)

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", line))
            _set_para_spacing(p, after=3)

        # Blank line — skip
        elif line.strip() == "":
            pass

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)
            _set_para_spacing(p, after=6)

        i += 1

    return doc


def _add_inline(para, text: str):
    """Parse inline markdown (bold, italic, links) and add runs to a paragraph."""
    # Pattern: **bold**, *italic*, [text](url)
    pattern = re.compile(
        r'\*\*(.+?)\*\*'       # bold
        r'|\*(.+?)\*'           # italic
        r'|\[(.+?)\]\((.+?)\)'  # link
    )
    pos = 0
    for m in pattern.finditer(text):
        # Text before match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])

        if m.group(1) is not None:  # bold
            run = para.add_run(m.group(1))
            run.bold = True
        elif m.group(2) is not None:  # italic
            run = para.add_run(m.group(2))
            run.italic = True
        elif m.group(3) is not None:  # link — keep as plain markdown text so URL survives round-trip
            para.add_run(f"[{m.group(3)}]({m.group(4)})")

        pos = m.end()

    # Remaining text
    if pos < len(text):
        para.add_run(text[pos:])


# ── Read back ────────────────────────────────────────────────────────────────

def read_docx(path: Path) -> tuple:
    """
    Read an edited .docx and return (title, description, body_markdown).
    Converts Word heading styles back to markdown headings.
    """
    doc = Document(path)

    title = ""
    description = ""
    body_lines = []
    in_body = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not in_body:
            if text.startswith("TITLE:"):
                title = text[6:].strip()
            elif text.startswith("DESC:"):
                description = text[5:].strip()
            elif text.startswith("─"):
                in_body = True
            continue

        style_name = para.style.name if para.style else ""

        if text in ("---", "─" * 60) or text.startswith("─"):
            continue  # skip dividers
        elif "Heading 2" in style_name:
            body_lines.append(f"\n## {text}")
        elif "Heading 3" in style_name:
            body_lines.append(f"\n### {text}")
        elif "List Bullet" in style_name:
            body_lines.append(f"- {_runs_to_markdown(para)}")
        elif "List Number" in style_name:
            body_lines.append(f"1. {_runs_to_markdown(para)}")
        elif text in ("Pros:", "Cons:"):
            body_lines.append(f"\n**{text}**")
        elif text:
            body_lines.append(_runs_to_markdown(para))
        else:
            body_lines.append("")

    body = "\n".join(body_lines).strip()
    return title, description, body


def _runs_to_markdown(para) -> str:
    """Convert paragraph runs back to inline markdown."""
    result = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        # Skip the grey URL annotations we added (font size 8)
        if run.font.size and run.font.size <= Pt(8):
            continue
        if run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        result.append(t)
    return "".join(result)
